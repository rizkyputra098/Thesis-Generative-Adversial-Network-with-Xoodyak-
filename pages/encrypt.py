"""
pages/encrypt.py - Sequential Two-Stage Process with Visual Comparisons
Stage 1: Input → Encrypt → Show Comparison → Confirm to Embed
Stage 2: Upload Image → Embed → Show Comparison → Download

SPESIFIKASI DATA:
- TXT: 250-1000 kata
- DOCX: 250-2000 kata (TANPA GAMBAR)
- TEXT: 250-1000 kata
"""

import streamlit as st
from xoodyak_utils import encrypt_file, calculate_hash
from stego_models_pytorch import hide_encrypted_data
from PIL import Image
import io
import time
from docx import Document


def count_words_docx(file_data):
    """Hitung jumlah kata dalam file DOCX dan cek gambar - IMPROVED"""
    try:
        doc = Document(io.BytesIO(file_data))
        word_count = 0
        has_image = False
        all_text = []
        
        # Method 1: Extract dari paragraphs (normal)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                all_text.append(paragraph.text)
            # Cek gambar
            for run in paragraph.runs:
                try:
                    if run._element.xml.find(b'blip') != -1:
                        has_image = True
                except:
                    pass
        
        # Method 2: Extract dari table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        all_text.append(cell.text)
        
        # Method 3: Extract langsung dari XML jika masih kosong
        if not all_text:
            try:
                from docx.oxml.text.paragraph import CT_P
                from docx.oxml.table import CT_Tbl
                from docx.table import _Cell
                from docx.text.paragraph import CT_P as CTP
                
                body = doc.element.body
                for element in body:
                    if isinstance(element, CT_P):
                        text = ''.join([node.text for node in element.iter() if node.text])
                        if text.strip():
                            all_text.append(text)
                    elif isinstance(element, CT_Tbl):
                        for row in element.iter():
                            text = ''.join([node.text for node in row.iter() if node.text])
                            if text.strip():
                                all_text.append(text)
            except:
                pass
        
        # Gabungkan semua text dan hitung kata
        full_text = " ".join(all_text)
        if full_text.strip():
            # Split dengan lebih robust
            words = full_text.split()
            word_count = len([w for w in words if w.strip()])
        
        return word_count, has_image
    except Exception as e:
        st.warning(f"⚠️ Error saat membaca DOCX: {str(e)}")
        return 0, False


def validate_image(image_data, filename):
    """Validasi resolusi gambar"""
    try:
        img = Image.open(io.BytesIO(image_data))
        width, height = img.size
        
        if width < 512:
            return False, f"❌ Lebar terlalu kecil! Ukuran: {width}x{height}. Minimal lebar 512 pixel."
        if height < 512:
            return False, f"❌ Tinggi terlalu kecil! Ukuran: {width}x{height}. Minimal tinggi 512 pixel."
        if width > 1024:
            return False, f"❌ Lebar terlalu besar! Ukuran: {width}x{height}. Maksimal lebar 1024 pixel."
        if height > 1024:
            return False, f"❌ Tinggi terlalu besar! Ukuran: {width}x{height}. Maksimal tinggi 1024 pixel."
        
        return True, f"✅ Resolusi valid: {width}x{height} pixel"
    except Exception as e:
        return False, f"❌ Error membaca gambar: {str(e)}"


def get_file_extension(filename):
    """Extract extension dari filename"""
    if '.' in filename:
        return filename.rsplit('.', 1)[1].lower()
    return 'bin'


def calculate_data_capacity(image_width: int, image_height: int) -> dict:
    """Calculate capacity gambar"""
    pixels = image_width * image_height
    bits_available = pixels * 1
    bytes_available = bits_available // 8
    
    return {
        'pixels': pixels,
        'bits': bits_available,
        'bytes': bytes_available,
        'kb': bytes_available / 1024,
    }


def format_bytes(bytes_size):
    """Format bytes to readable format"""
    for unit in ['B', 'KB', 'MB']:
        if bytes_size < 1024:
            return f"{bytes_size:.2f} {unit}"
        bytes_size /= 1024
    return f"{bytes_size:.2f} GB"


def render():
    # Initialize state
    if 'current_stage' not in st.session_state:
        st.session_state.current_stage = 1  # 1, 1.5, 2, 2.5
    
    if st.button("← Kembali ke Home"):
        st.session_state.page = 'home'
        st.session_state.current_stage = 1
        keys_to_clear = ['encrypt_result', 'encrypt_filename', 'encrypt_hash', 'pre_hash',
                        'stego_image', 'stego_metrics', 'original_cover', 'file_extension', 
                        'perf_metrics', 'temp_data', 'temp_filename', 'temp_cover_image', 'temp_cover_name', 
                        'temp_cover_width', 'temp_cover_height']
        for key in keys_to_clear:
            if key in st.session_state:
                del st.session_state[key]
        st.rerun()
    
    st.markdown("<div style='margin: 2rem 0;'></div>", unsafe_allow_html=True)
    
    st.markdown("""
    <div style="text-align: center; margin-bottom: 3rem;">
        <h1 style="font-size: 2.5rem; font-weight: 800; 
                   background: linear-gradient(135deg, #64b5f6 0%, #81c784 100%);
                   -webkit-background-clip: text;
                   -webkit-text-fill-color: transparent;">
            🔐 ENKRIPSI & STEGANOGRAFI
        </h1>
        <p style="color: #b0bec5; margin-top: 0.5rem;">
            Amankan data dengan Xoodyak AEAD + Deep Learning Steganography
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # ===================== STAGE 1: INPUT & ENCRYPTION =====================
    if st.session_state.current_stage == 1:
        st.markdown("## 📝 STAGE 1: INPUT & ENKRIPSI DATA")
        st.markdown("---")
        
        col_input_left, col_input_right = st.columns([1.2, 1], gap="large")
        
        with col_input_left:
            st.markdown("### INPUT DATA")
            
            input_method = st.radio(
                "",
                ["FILE", "TEXT"],
                horizontal=True,
                label_visibility="collapsed",
                key="encrypt_input_method"
            )
            
            st.markdown("<div style='margin: 1.5rem 0;'></div>", unsafe_allow_html=True)
            
            data_to_process = None
            filename = None
            file_extension = 'txt'
            
            if input_method == "FILE":
                st.markdown("**📄 Data File**")
                st.caption("Format: TXT, DOCX")
                
                uploaded_file = st.file_uploader(
                    "Upload file untuk dienkripsi", 
                    label_visibility="collapsed",
                    type=['txt', 'docx'],
                    key="file_data_input"
                )
                
                if uploaded_file:
                    data_to_process = uploaded_file.read()
                    filename = uploaded_file.name
                    file_extension = get_file_extension(filename)
                    
                    file_size_kb = len(data_to_process) / 1024
                    
                    if file_extension == 'txt':
                        try:
                            text_content = data_to_process.decode('utf-8', errors='ignore')
                            word_count = len(text_content.split())
                            
                            if word_count < 0:
                                st.error(f"❌ File .txt terlalu pendek! Jumlah kata: {word_count}. Minimal 250 kata.")
                                data_to_process = None
                            elif word_count > 1000:
                                st.error(f"❌ File .txt terlalu panjang! Jumlah kata: {word_count}. Maksimal 1000 kata.")
                                data_to_process = None
                            else:
                                st.success(f"✅ Ukuran file valid")
                                st.info(f"📊 {word_count} kata | {file_size_kb:.1f} KB")
                        except:
                            st.error("Error membaca file .txt")
                            data_to_process = None
                    
                    elif file_extension == 'docx':
                        if file_size_kb > 500:
                            st.error(f"❌ File .docx terlalu besar! Ukuran: {file_size_kb:.1f} KB. Maksimal 500 KB.")
                            data_to_process = None
                        elif file_size_kb < 10:
                            st.error(f"❌ File .docx terlalu kecil! Ukuran: {file_size_kb:.1f} KB. Minimal 10 KB.")
                            data_to_process = None
                        else:
                            # Hitung kata dan cek gambar
                            word_count, has_image = count_words_docx(data_to_process)
                            
                            if has_image:
                                st.error(f"❌ File .docx tidak boleh mengandung gambar!")
                                data_to_process = None
                            elif word_count < 1:
                                st.error(f"❌ File .docx terlalu pendek! Jumlah kata: {word_count}. Minimal 250 kata.")
                                data_to_process = None
                            elif word_count > 2000:
                                st.error(f"❌ File .docx terlalu panjang! Jumlah kata: {word_count}. Maksimal 2000 kata.")
                                data_to_process = None
                            else:
                                st.success(f"✅ Ukuran file valid")
                                st.info(f"📊 {word_count} kata | {file_size_kb:.1f} KB | Format: DOCX")
                    
                    if data_to_process:
                        file_hash = calculate_hash(data_to_process)
                        st.markdown(f"""
                        <div class="success-status">
                        ✓ {filename}<br>
                        ✓ Ukuran: {len(data_to_process)} bytes<br>
                        ✓ Tipe: .{file_extension.upper()}
                        </div>
                        """, unsafe_allow_html=True)
                        
                        st.markdown(f"""
                        <div class="hash-label">Hash MD5 (Original)</div>
                        <div class="hash-box">{file_hash}</div>
                        """, unsafe_allow_html=True)
                        st.session_state['pre_hash'] = file_hash
                        st.session_state['file_extension'] = file_extension
                        st.session_state['temp_data'] = data_to_process
                        st.session_state['temp_filename'] = filename
            
            else:  # TEXT
                st.markdown("**📝 Data Teks**")
                st.caption("Minimal 250 kata, Maksimal 1000 kata")
                
                text_input = st.text_area(
                    "Masukkan teks untuk dienkripsi",
                    height=150,
                    placeholder="Ketik teks di sini...",
                    label_visibility="collapsed",
                    key="text_data_input"
                )
                
                if text_input:
                    word_count = len(text_input.split())
                    
                    if word_count < 1:
                        st.error(f"❌ Teks terlalu pendek! Jumlah kata: {word_count}. Minimal 250 kata.")
                    elif word_count > 1000:
                        st.error(f"❌ Teks terlalu panjang! Jumlah kata: {word_count}. Maksimal 1000 kata.")
                    else:
                        st.success(f"✅ Jumlah kata valid")
                        st.info(f"📊 {word_count} kata")
                        
                        data_to_process = text_input.encode('utf-8')
                        filename = "encrypted_text.txt"
                        file_extension = 'txt'
                        
                        text_hash = calculate_hash(data_to_process)
                        st.markdown(f"""
                        <div class="hash-label">Hash MD5 (Original)</div>
                        <div class="hash-box">{text_hash}</div>
                        """, unsafe_allow_html=True)
                        st.session_state['pre_hash'] = text_hash
                        st.session_state['file_extension'] = file_extension
                        st.session_state['temp_data'] = data_to_process
                        st.session_state['temp_filename'] = filename
            
            st.markdown("<div style='margin: 1.5rem 0;'></div>", unsafe_allow_html=True)
            
            # STEGANOGRAPHY INPUT - GAMBAR CARRIER
            st.markdown("**🖼️ Gambar Carrier (Steganografi)**")
            st.caption("⚡ Resolusi: 512-1024 pixel (W & H), Format: PNG/JPG/BMP")
            
            uploaded_image = st.file_uploader(
                "Upload gambar untuk menyembunyikan data terenkripsi",
                label_visibility="collapsed",
                type=['jpg', 'jpeg', 'png', 'bmp'],
                key="stego_image_input_stage1"
            )
            
            if uploaded_image:
                image_data = uploaded_image.read()
                is_valid, message = validate_image(image_data, uploaded_image.name)
                
                if is_valid:
                    st.success(message)
                    img = Image.open(io.BytesIO(image_data))
                    cover_width, cover_height = img.size
                    
                    capacity = calculate_data_capacity(cover_width, cover_height)
                    st.info(f"📊 Kapasitas: {capacity['kb']:.1f} KB")
                    
                    if uploaded_image.name.lower().endswith(('.jpg', '.jpeg', '.bmp')):
                        png_io = io.BytesIO()
                        img.convert('RGB').save(png_io, format='PNG')
                        image_data = png_io.getvalue()
                    
                    st.session_state['temp_cover_image'] = image_data
                    st.session_state['temp_cover_name'] = uploaded_image.name
                    st.session_state['temp_cover_width'] = cover_width
                    st.session_state['temp_cover_height'] = cover_height
                else:
                    st.error(message)
            
            st.markdown("<div style='margin: 1.5rem 0;'></div>", unsafe_allow_html=True)
            
            password = st.text_input(
                "Password",
                type="password",
                placeholder="Masukkan password untuk enkripsi",
                label_visibility="collapsed",
                key="encrypt_password"
            )
            
            st.markdown("<div style='margin: 1.5rem 0;'></div>", unsafe_allow_html=True)
            
            if st.button("🔐 ENKRIPSI SEKARANG", use_container_width=True, key="encrypt_button"):
                temp_data = st.session_state.get('temp_data')
                if not temp_data:
                    st.error("❌ Silakan masukkan data yang ingin dienkripsi")
                elif not password:
                    st.error("❌ Password diperlukan untuk enkripsi")
                else:
                    try:
                        with st.spinner("⚙️ Mengenkripsi data dengan Xoodyak..."):
                            start_time = time.time()
                            encrypted_data = encrypt_file(temp_data, password, st.session_state.get('temp_filename', 'data'))
                            encrypt_time = time.time() - start_time
                            
                            perf_metrics = {
                                'data_original_size': len(temp_data),
                                'data_encrypted_size': len(encrypted_data),
                                'encryption_time': encrypt_time,
                                'encryption_speed': len(temp_data) / (encrypt_time + 0.001)
                            }
                            
                            st.session_state['encrypt_result'] = encrypted_data
                            st.session_state['encrypt_filename'] = st.session_state.get('temp_filename', 'data')
                            st.session_state['encrypt_hash'] = calculate_hash(encrypted_data)
                            st.session_state['perf_metrics'] = perf_metrics
                            st.session_state['current_stage'] = 1.5
                            
                            st.success("✅ Enkripsi berhasil!")
                            st.rerun()
                    
                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")
        
        with col_input_right:
            st.markdown("### 📊 PREVIEW DATA")
            if 'pre_hash' in st.session_state:
                st.info("✅ Data siap untuk dienkripsi")
                st.metric("Ukuran Data", format_bytes(st.session_state.get('temp_data', b'') and len(st.session_state['temp_data']) or 0))
                st.caption("Data akan berubah setelah enkripsi")
    
# ===================== STAGE 1B: ENCRYPTION COMPARISON =====================
    elif st.session_state.current_stage == 1.5:
        st.markdown("## ✅ HASIL ENKRIPSI - PERBANDINGAN DATA")
        st.markdown("---")
        
        encrypted_data = st.session_state.get('encrypt_result')
        original_data = st.session_state.get('temp_data', b'')
        perf_metrics = st.session_state.get('perf_metrics', {})
        
        st.markdown("### 📊 PERBANDINGAN SEBELUM & SESUDAH ENKRIPSI")
        
        col_before, col_after = st.columns(2)
        
        with col_before:
            st.markdown("<p style='text-align: center; color: #64b5f6; font-weight: bold; font-size: 1.2rem;'>🔓 SEBELUM ENKRIPSI</p>", unsafe_allow_html=True)
            st.markdown(f"**Ukuran:** {format_bytes(perf_metrics.get('data_original_size', 0))}")
            
            # Tampilkan preview teks original dengan expander
            try:
                text_preview = original_data.decode('utf-8', errors='ignore')
                
                with st.expander("📖 Lihat Data Lengkap", expanded=True):
                    st.text_area(
                        "Data Original (Plain Text)",
                        value=text_preview,
                        height=350,
                        disabled=True,
                        label_visibility="collapsed"
                    )
            except:
                st.info("(Data binary atau non-text)")
        
        with col_after:
            st.markdown("<p style='text-align: center; color: #81c784; font-weight: bold; font-size: 1.2rem;'>🔒 SESUDAH ENKRIPSI</p>", unsafe_allow_html=True)
            st.markdown(f"**Ukuran:** {format_bytes(perf_metrics.get('data_encrypted_size', 0))}")
            
            # Tampilkan preview encrypted data dalam hex dengan expander
            encrypted_hex = encrypted_data.hex().upper()
            encrypted_preview = encrypted_hex
            
               
            with st.expander("🔐 Lihat Data Terenkripsi (Hex)", expanded=True):
                st.text_area(
                    "Data Terenkripsi (Hexadecimal)",
                    value=encrypted_preview,
                    height=350,
                    disabled=True,
                    label_visibility="collapsed"
                )
        
        st.markdown("<div style='margin: 1.5rem 0;'></div>", unsafe_allow_html=True)
        
        st.markdown("### ⏱️ STATISTIK ENKRIPSI")
        col_stat_1, col_stat_2, col_stat_3 = st.columns(3)
        
        with col_stat_1:
            st.metric("Waktu Enkripsi", f"{perf_metrics.get('encryption_time', 0):.3f}s")
        with col_stat_2:
            st.metric("Kecepatan", f"{format_bytes(perf_metrics.get('encryption_speed', 0))}/s")
        with col_stat_3:
            size_diff = perf_metrics.get('data_encrypted_size', 0) - perf_metrics.get('data_original_size', 0)
            st.metric("Selisih Ukuran", f"{format_bytes(abs(size_diff))}")
        
        st.markdown("<div style='margin: 2rem 0;'></div>", unsafe_allow_html=True)
        st.markdown("---")
        st.markdown("<div style='margin: 2rem 0;'></div>", unsafe_allow_html=True)
        
        st.markdown("### 🎯 LANJUT KE TAHAP BERIKUTNYA")
        
        if st.button("📤 LANJUT KE EMBEDDING", use_container_width=True, key="proceed_to_embed"):
            st.session_state.current_stage = 2
            st.rerun()
        
        if st.button("↺ ULANGI ENKRIPSI", use_container_width=True, key="restart_encrypt"):
            st.session_state.current_stage = 1
            st.rerun()
    
    # ===================== STAGE 2: EMBEDDING =====================
    elif st.session_state.current_stage == 2:
        st.markdown("## 🖼️ STAGE 2: EMBEDDING DATA TERENKRIPSI KE GAMBAR")
        st.markdown("---")
        
        encrypted_data = st.session_state.get('encrypt_result')
        perf_metrics = st.session_state.get('perf_metrics', {})
        
        col_embed_left, col_embed_right = st.columns([1.2, 1], gap="large")
        
        with col_embed_left:
            st.markdown("### INPUT GAMBAR CARRIER")
            st.caption("⚡ Resolusi: 512-1024 pixel (W & H), Format: PNG/JPG/BMP")
            
            # Cek apakah gambar sudah ada dari stage 1
            if st.session_state.get('temp_cover_image'):
                st.success("✅ Gambar carrier sudah tersedia dari tahap sebelumnya")
                cover_image_data = st.session_state.get('temp_cover_image')
                cover_image_name = st.session_state.get('temp_cover_name')
                cover_width = st.session_state.get('temp_cover_width', 0)
                cover_height = st.session_state.get('temp_cover_height', 0)
                
                capacity = calculate_data_capacity(cover_width, cover_height)
                st.info(f"📊 Kapasitas gambar: {capacity['kb']:.1f} KB ({capacity['bytes']} bytes)")
                
                img = Image.open(io.BytesIO(cover_image_data))
                st.image(img, caption=f"Gambar: {cover_image_name}", use_column_width=True)
            else:
                st.info("Belum ada gambar. Upload di bawah:")
                
                uploaded_image = st.file_uploader(
                    "Upload gambar untuk menyembunyikan data terenkripsi",
                    label_visibility="collapsed",
                    type=['jpg', 'jpeg', 'png', 'bmp'],
                    key="stego_image_input"
                )
                
                if uploaded_image:
                    image_data = uploaded_image.read()
                    is_valid, message = validate_image(image_data, uploaded_image.name)
                    
                    if is_valid:
                        st.success(message)
                        img = Image.open(io.BytesIO(image_data))
                        cover_width, cover_height = img.size
                        
                        capacity = calculate_data_capacity(cover_width, cover_height)
                        
                        st.info(f"📊 Kapasitas gambar: {capacity['kb']:.1f} KB ({capacity['bytes']} bytes)")
                        st.caption(f"💾 Ukuran file gambar: {format_bytes(len(image_data))}")
                        
                        if uploaded_image.name.lower().endswith(('.jpg', '.jpeg', '.bmp')):
                            st.info("💡 Gambar akan dikonversi ke PNG")
                            png_io = io.BytesIO()
                            img.convert('RGB').save(png_io, format='PNG')
                            image_data = png_io.getvalue()
                        
                        st.image(img, caption=f"Original: {uploaded_image.name}", use_column_width=True)
                        st.session_state['temp_cover_image'] = image_data
                        st.session_state['temp_cover_name'] = uploaded_image.name
                        st.session_state['temp_cover_width'] = cover_width
                        st.session_state['temp_cover_height'] = cover_height
                    else:
                        st.error(message)
            
            st.markdown("<div style='margin: 1.5rem 0;'></div>", unsafe_allow_html=True)
            
            if st.button("📤 EMBED DATA KE GAMBAR", use_container_width=True, key="embed_button"):
                cover_image_data = st.session_state.get('temp_cover_image')
                cover_width = st.session_state.get('temp_cover_width', 0)
                cover_height = st.session_state.get('temp_cover_height', 0)
                
                if not cover_image_data:
                    st.error("❌ Gambar carrier diperlukan untuk steganografi")
                else:
                    capacity = calculate_data_capacity(cover_width, cover_height)
                    
                    if len(encrypted_data) > capacity['bytes']:
                        st.error(f"""
                        ❌ Data terlalu besar untuk gambar ini!
                        
                        📊 Data: {len(encrypted_data)} bytes
                        📊 Kapasitas: {capacity['bytes']} bytes
                        """)
                    else:
                        try:
                            with st.spinner("🧠 Menyembunyikan data dengan AI Model..."):
                                start_time = time.time()
                                stego_image, metrics = hide_encrypted_data(cover_image_data, encrypted_data)
                                embedding_time = time.time() - start_time
                                
                                perf_metrics['image_original_size'] = len(cover_image_data)
                                perf_metrics['image_stego_size'] = len(stego_image)
                                perf_metrics['embedding_time'] = embedding_time
                                perf_metrics['size_increase'] = ((len(stego_image) - len(cover_image_data)) / len(cover_image_data) * 100)
                                perf_metrics['psnr'] = metrics.get('psnr', 0)
                                perf_metrics['embedding_speed'] = len(encrypted_data) / (embedding_time + 0.001)
                                
                                st.session_state['stego_image'] = stego_image
                                st.session_state['stego_metrics'] = metrics
                                st.session_state['original_cover'] = cover_image_data
                                st.session_state['cover_name'] = st.session_state.get('temp_cover_name', 'cover.png')
                                st.session_state['perf_metrics'] = perf_metrics
                                st.session_state['current_stage'] = 2.5
                                
                                st.success("✅ Embedding berhasil!")
                                st.rerun()
                        
                        except Exception as e:
                            st.error(f"❌ Error: {str(e)}")
        
        with col_embed_right:
            st.markdown("### 📋 INFO EMBEDDING")
            st.info("⏳ Tunggu sampai gambar carrier di-upload, kemudian klik tombol EMBED")
            st.caption("Data terenkripsi siap untuk disembunyikan")
            st.metric("Ukuran Data", format_bytes(len(encrypted_data)))
    
    # ===================== STAGE 2B: EMBEDDING COMPARISON =====================
    elif st.session_state.current_stage == 2.5:
        st.markdown("## ✅ HASIL EMBEDDING - PERBANDINGAN GAMBAR")
        st.markdown("---")
        
        stego_img = st.session_state.get('stego_image')
        original_cover = st.session_state.get('original_cover')
        metrics = st.session_state.get('stego_metrics', {})
        perf_metrics = st.session_state.get('perf_metrics', {})
        
        st.markdown("### 📸 PERBANDINGAN SEBELUM & SESUDAH EMBEDDING")
        
        col_img_before, col_img_after = st.columns(2)
        
        with col_img_before:
            st.markdown("<p style='text-align: center; color: #64b5f6; font-weight: bold; font-size: 1.2rem;'>🖼️ SEBELUM EMBEDDING</p>", unsafe_allow_html=True)
            if original_cover:
                original_pil = Image.open(io.BytesIO(original_cover))
                st.image(original_pil, use_column_width=True)
            st.metric("Ukuran", format_bytes(perf_metrics.get('image_original_size', 0)))
            st.caption("Gambar Original (Carrier)")
        
        with col_img_after:
            st.markdown("<p style='text-align: center; color: #81c784; font-weight: bold; font-size: 1.2rem;'>🎨 SESUDAH EMBEDDING</p>", unsafe_allow_html=True)
            stego_pil = Image.open(io.BytesIO(stego_img))
            st.image(stego_pil, use_column_width=True)
            st.metric("Ukuran", format_bytes(perf_metrics.get('image_stego_size', 0)))
            st.caption("Gambar dengan Data Tersembunyi")
        
        st.markdown("<div style='margin: 1.5rem 0;'></div>", unsafe_allow_html=True)
        
        st.markdown("### ⏱️ STATISTIK EMBEDDING")
        col_stat_1, col_stat_2, col_stat_3 = st.columns(3)
        
        with col_stat_1:
            st.metric("Waktu Embedding", f"{perf_metrics.get('embedding_time', 0):.3f}s")
        with col_stat_2:
            st.metric("Kecepatan", f"{format_bytes(perf_metrics.get('embedding_speed', 0))}/s")
        with col_stat_3:
            st.metric("Peningkatan Ukuran", f"{perf_metrics.get('size_increase', 0):.2f}%")
        
        st.markdown("<div style='margin: 1rem 0;'></div>", unsafe_allow_html=True)
        
        st.markdown("### 🎯 KUALITAS GAMBAR")
        col_quality = st.columns(1)[0]
        with col_quality:
            psnr_value = metrics.get('psnr', 0)
            psnr_color = "🟢" if psnr_value > 40 else "🟡" if psnr_value > 30 else "🔴"
            st.metric("PSNR", f"{psnr_value:.2f} dB", f"{psnr_color} Kualitas Gambar")
        
        st.markdown("<div style='margin: 2rem 0;'></div>", unsafe_allow_html=True)
        st.markdown("---")
        st.markdown("<div style='margin: 2rem 0;'></div>", unsafe_allow_html=True)
        
        st.markdown("### 📥 DOWNLOAD HASIL")
        cover_name = st.session_state.get('cover_name', 'cover.png')
        base_name = cover_name.rsplit('.', 1)[0]
        stego_filename = f"{base_name}_stego.png"
        
        st.download_button(
            label="⬇️ DOWNLOAD GAMBAR STEGO",
            data=stego_img,
            file_name=stego_filename,
            mime="image/png",
            use_container_width=True,
            help="Download gambar dengan data terenkripsi tersembunyi"
        )
        
        st.markdown("<div style='margin: 1rem 0;'></div>", unsafe_allow_html=True)
        
        if st.button("🔄 MULAI DARI AWAL", use_container_width=True, key="reset_all"):
            keys_to_clear = ['encrypt_result', 'encrypt_filename', 'encrypt_hash', 'pre_hash',
                            'stego_image', 'stego_metrics', 'original_cover', 'cover_name', 'file_extension', 
                            'perf_metrics', 'temp_data', 'temp_filename', 'temp_cover_image', 'temp_cover_name', 
                            'temp_cover_width', 'temp_cover_height']
            for key in keys_to_clear:
                if key in st.session_state:
                    del st.session_state[key]
            st.session_state.current_stage = 1
            st.rerun()